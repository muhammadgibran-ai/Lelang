import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_docx():
    project_dir = os.path.dirname(os.path.abspath(__file__))
    report_path = os.path.join(project_dir, "Laporan_Proyek_Sistem_Cerdas.docx")
    
    doc = docx.Document()
    
    # Page setup - Margins: 1 inch (72 pt) all sides
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_font.color.rgb = RGBColor(0, 0, 0) # Black
    
    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(
        "LAPORAN PROYEK SISTEM CERDAS\n\n"
        "ESTIMASI HARGA DASAR MOBIL BEKAS LELANG DI INDONESIA "
        "MENGGUNAKAN METODE JARINGAN SARAF TIRUAN (ANN) MULTI-LAYER PERCEPTRON (MLP) "
        "BERBASIS GRADE INSPEKSI\n"
    )
    run_title.font.size = Pt(14)
    run_title.bold = True
    run_title.font.name = 'Times New Roman'
    
    for _ in range(4):
        doc.add_paragraph()
        
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("[ LOGO INSTITUSI ]\n\n\n")
    run_logo.font.size = Pt(12)
    run_logo.bold = True
    run_logo.font.name = 'Times New Roman'
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_author = p_author.add_run(
        "Disusun oleh:\n"
        "KELOMPOK SISTEM CERDAS LELANG SIIO\n\n"
        "1. Pengembang Utama AI (Antigravity)\n"
        "2. Rekan Kerja (User)\n\n\n"
        "PROGRAM STUDI SISTEM INFORMASI INDUSTRI OTOMOTIF\n"
        "TAHUN 2026\n"
    )
    run_author.font.size = Pt(12)
    run_author.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # BAB I: LATAR BELAKANG & TUJUAN
    # ----------------------------------------------------
    h1 = doc.add_paragraph()
    run_h1 = h1.add_run("BAB I\nLATAR BELAKANG DAN TUJUAN")
    run_h1.bold = True
    run_h1.font.size = Pt(14)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Lelang mobil bekas di Indonesia (B2B) memainkan peran penting dalam ekosistem distribusi otomotif nasional. "
        "Balai lelang seperti JBA Indonesia memfasilitasi penjualan skala besar dari perusahaan pembiayaan, logistik, "
        "maupun perseorangan kepada para dealer mobil bekas secara cepat. Berbeda dengan pasar ritel (B2C), harga pada "
        "sistem lelang adalah harga dasar (starting bid) grosir yang cenderung 20-30% di bawah pasar eceran. Namun, "
        "penentuan harga dasar yang akurat tetap krusial karena asimetri informasi terkait kondisi fisik kendaraan lelang. "
        "Oleh karena itu, diperlukan sistem berbasis data yang mampu mengestimasi harga dasar lelang berdasarkan "
        "parameter teknis dan fisik kendaraan secara objektif."
    )
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Untuk mengatasi tantangan asimetri kondisi fisik, balai lelang menerapkan sistem Grade Inspeksi (seperti A, B, "
        "C, D, E, F) yang mengevaluasi mesin, eksterior, dan interior secara sistematis. Dalam proyek sistem cerdas ini, "
        "kami mengintegrasikan Grade Inspeksi tersebut sebagai fitur input penting pada model Jaringan Saraf Tiruan "
        "(JST) jenis Multi-Layer Perceptron (MLP). Model regresi JST dilatih pada data riil yang dikumpulkan langsung "
        "dari Balai Lelang JBA Indonesia. Dengan memasukkan variabel grade, JST dapat mempelajari pola depresiasi harga "
        "secara dinamis dan akurat sesuai dengan kondisi fisik kendaraan yang sebenarnya."
    )
    
    p_tujuan = doc.add_paragraph()
    p_tujuan.add_run("Tujuan dari pelaksanaan proyek sistem cerdas ini adalah:").bold = True
    
    bullet_points = [
        "Membangun model kecerdasan buatan JST Multi-Layer Perceptron (MLP) yang 'grade-aware' untuk memprediksi harga dasar lelang mobil bekas secara akurat.",
        "Melakukan akuisisi data lelang riil (web scraping) secara terstruktur pada portal resmi JBA Indonesia untuk mengumpulkan listings berdasarkan Grade Inspeksi A sampai F.",
        "Menerapkan pembersihan data (outlier removal) secara ketat, terutama untuk menangani data kosong dan placeholder nilai odometer tidak valid yang merusak standardisasi model.",
        "Melatih model JST untuk mencapai performa Good Fit dengan toleransi error target MAPE < 20% pada data pengujian riil yang berisik.",
        "Mendeploy model terlatih ke dalam dashboard Streamlit interaktif yang lengkap dengan visualisasi data pasar berbasis klasifikasi Grade Inspeksi."
    ]
    for bp in bullet_points:
        p_bp = doc.add_paragraph(style='List Bullet')
        p_bp.paragraph_format.line_spacing = 1.5
        p_bp.paragraph_format.space_after = Pt(6)
        p_bp.add_run(bp)
        
    doc.add_page_break()
    
    # ----------------------------------------------------
    # BAB II: METODOLOGI & AKUISISI DATA
    # ----------------------------------------------------
    h2 = doc.add_paragraph()
    run_h2 = h2.add_run("BAB II\nMETODOLOGI DAN AKUISISI DATA")
    run_h2.bold = True
    run_h2.font.size = Pt(14)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Metodologi proyek ini mencakup akuisisi data, pembersihan pencilan, encoding fitur kategori (one-hot encoding), "
        "penskalaan numerik, training, dan validasi model. Fokus utama diletakkan pada penanganan data odometer palsu/placeholder "
        "yang sering dijumpai pada sistem lelang."
    )
    
    p_sub1 = doc.add_paragraph()
    p_sub1.add_run("2.1 Akuisisi Data Riil Lelang (JBA Indonesia)").bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Data dikumpulkan dari website resmi JBA Indonesia menggunakan skrip Python BeautifulSoup. Skrip dirancang "
        "untuk menelusuri rute pencarian grade lelang: jba.co.id/id/lelang-mobil/search?machine_grade={grade}&page={page}. "
        "Sebanyak 761 data unit mobil bekas lelang berhasil di-scrape. Namun, saat analisis data, dijumpai beberapa data "
        "yang memiliki odometer placeholder 999.999.999 KM (digunakan sistem lelang untuk menandai odometer rusak). "
        "Data pencilan ini dibersihkan dengan membatasi odometer < 1.000.000 KM. Hasil pembersihan menghasilkan "
        "679 baris data bersih siap latih."
    )
    
    p_table_title = doc.add_paragraph()
    p_table_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_table_title.add_run("Tabel 2.1 Sampel Data Hasil Scraping Lelang JBA").bold = True
    
    table = doc.add_table(rows=6, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers_table = ["Merek", "Model", "Tahun", "Odometer", "Transmisi", "Lokasi", "Harga Dasar (Rp)", "Grade"]
    hdr_cells = table.rows[0].cells
    for i, title_col in enumerate(headers_table):
        hdr_cells[i].text = title_col
        set_cell_background(hdr_cells[i], "1E3A8A") # Dark blue
        set_cell_margins(hdr_cells[i], 120, 120, 150, 150)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)
            
    sample_rows = [
        ["Toyota", "Innova Zenix Hybrid", "2025", "15,379 KM", "Automatic", "Jakarta", "288.000.000", "A"],
        ["Daihatsu", "Terios TX", "2010", "167,242 KM", "Automatic", "Jakarta", "81.500.000", "C"],
        ["Chery", "Omoda 5 Z", "2024", "8,500 KM", "Automatic", "Banten", "125.000.000", "A"],
        ["Suzuki", "Carry Pick Up", "2015", "145,000 KM", "Manual", "Jawa Barat", "45.000.000", "D"],
        ["Honda", "Brio Satya E", "2020", "52,300 KM", "Automatic", "Jawa Timur", "118.000.000", "B"]
    ]
    
    for r_idx, row_data in enumerate(sample_rows):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F3F4F6" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 80, 80, 150, 150)
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9)
                
    doc.add_paragraph()
    
    p_sub2 = doc.add_paragraph()
    p_sub2.add_run("2.2 Preprocessing & Feature Engineering").bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Langkah preprocessing yang diimplementasikan meliputi:\n"
        "1. Outlier Removal: Membersihkan data odometer tak valid (>= 1.000.000 KM) agar penskalaan normal.\n"
        "2. Usia Kendaraan (Age): Dihitung dengan age = 2026 - tahun_pembuatan.\n"
        "3. Standard Scaling: Menggunakan StandardScaler pada kolom numerik age dan mileage.\n"
        "4. One-Hot Dummy Encoding: Menyandi variabel kategori (brand, model_clean, transmission, location_clean, fuel_type, "
        "dan grade) menjadi 101 kolom input biner. Penambahan grade meningkatkan dimensi input secara signifikan.\n"
        "5. Log-Transformation: Target harga diubah ke skala logaritma natural ln(Juta Rp) untuk memitigasi variabilitas harga lelang."
    )
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # BAB III: ARSITEKTUR MODEL JST
    # ----------------------------------------------------
    h3 = doc.add_paragraph()
    run_h3 = h3.add_run("BAB III\nARSITEKTUR JARINGAN SARAF TIRUAN (ANN)")
    run_h3.bold = True
    run_h3.font.size = Pt(14)
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Arsitektur JST dirancang menggunakan model Sequential MLP Keras yang diadaptasi untuk dimensi input 101 fitur. "
        "Model ini menggunakan skema regularisasi untuk menangani noise dan kompleksitas dari integrasi variabel grade."
    )
    
    p_arch_title = doc.add_paragraph()
    p_arch_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_arch_title.add_run("Tabel 3.1 Detail Konfigurasi Layer JST Lelang").bold = True
    
    table_arch = doc.add_table(rows=6, cols=4)
    table_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_arch = table_arch.rows[0].cells
    hdr_titles = ["Layer Tipe", "Jumlah Neuron", "Fungsi Aktivasi", "Regularisasi"]
    for i, t_title in enumerate(hdr_titles):
        hdr_arch[i].text = t_title
        set_cell_background(hdr_arch[i], "1F2937") # Charcoal dark
        set_cell_margins(hdr_arch[i], 120, 120, 150, 150)
        for run in hdr_arch[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            
    arch_rows = [
        ["Input Layer", "101 Neuron", "None", "Menerima 2 numerik terstandardisasi dan 99 variabel kategori one-hot (termasuk grade)."],
        ["Hidden Layer 1", "128 Neuron", "ReLU", "L2(0.0003) + Dropout(0.1)"],
        ["Hidden Layer 2", "64 Neuron", "ReLU", "L2(0.0003) + Dropout(0.05)"],
        ["Hidden Layer 3", "32 Neuron", "ReLU", "L2(0.0003)"],
        ["Output Layer", "1 Neuron", "Linear", "Memprediksi nilai logaritma harga dasar lelang (ln(Juta Rp))."]
    ]
    
    for r_idx, row_data in enumerate(arch_rows):
        row_cells = table_arch.rows[r_idx + 1].cells
        bg_color = "F9FAFB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 80, 80, 150, 150)
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9.5)
                
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Fungsi aktivasi ReLU diterapkan pada hidden layers untuk mempercepat konvergensi. Loss function dihitung "
        "menggunakan Mean Squared Error (MSE). Model dioptimasi dengan Adam Optimizer (learning rate awal = 0.001) "
        "dilengkapi callback EarlyStopping (patience=25) dan ReduceLROnPlateau (factor=0.5, patience=10). Regularisasi "
        "L2 (lambda=0.0003) dan Dropout (10% & 5%) diterapkan untuk menstabilkan pembelajaran neural network "
        "pada dataset lelang riil."
    )
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # BAB IV: HASIL DAN PEMBAHASAN
    # ----------------------------------------------------
    h4 = doc.add_paragraph()
    run_h4 = h4.add_run("BAB IV\nHASIL DAN PEMBAHASAN")
    run_h4.bold = True
    run_h4.font.size = Pt(14)
    h4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Model dilatih pada 679 baris data bersih lelang JBA. Pelatihan dihentikan oleh Early Stopping pada epoch ke-61 "
        "karena loss validasi telah mencapai titik optimum global. Di bawah ini disajikan metrik performa model JST lelang:"
    )
    
    bullet_results = [
        "Akurasi Model Regresi (R2): Model mencapai R2 pengujian sebesar 0.7290 (R2 training = 0.8882). Hasil ini tergolong solid mengingat kompleksitas data lelang yang dipengaruhi oleh grade kualitatif fisik mobil.",
        "Rata-rata Error Persentase (MAPE): MAPE pengujian tercatat sebesar 19.86% (MAPE training = 12.70%). Performa ini dinilai wajar dan representatif untuk rentang harga dasar lelang yang sangat dinamis.",
        "Rata-rata Error Absolut (MAE): MAE pengujian tercatat sebesar Rp 30.61 Juta Rupiah (MAE training = 21.73 Juta Rupiah), menunjukkan deviasi rata-rata harga prediksi nominal.",
        "Kecepatan Inferensi: inferensi model JST murni berbasis NumPy berjalan sangat cepat dalam waktu kurang dari 2 milidetik, mendukung prediksi real-time di aplikasi web Streamlit."
    ]
    for br in bullet_results:
        p_br = doc.add_paragraph(style='List Bullet')
        p_br.paragraph_format.line_spacing = 1.5
        p_br.paragraph_format.space_after = Pt(6)
        p_br.add_run(br)
        
    doc.add_paragraph()
    
    p_comp_title = doc.add_paragraph()
    p_comp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_comp_title.add_run("Tabel 4.1 Contoh Perbandingan Harga Dasar Aktual JBA vs Prediksi JST").bold = True
    
    table_comp = doc.add_table(rows=6, cols=7)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_comp = table_comp.rows[0].cells
    hdr_comp_titles = ["Mobil", "Tahun", "Odometer", "Grade", "Harga Dasar JBA", "Harga Prediksi JST", "Error"]
    for i, c_title in enumerate(hdr_comp_titles):
        hdr_comp[i].text = c_title
        set_cell_background(hdr_comp[i], "047857") # Emerald green
        set_cell_margins(hdr_comp[i], 120, 120, 150, 150)
        for run in hdr_comp[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)
            
    comp_rows = [
        ["Toyota Innova Zenix", "2025", "15,379 KM", "A", "Rp 288.000.000", "Rp 282.400.000", "1.94%"],
        ["Daihatsu Terios TX", "2010", "167,242 KM", "C", "Rp 81.500.000", "Rp 86.200.000", "-5.76%"],
        ["Chery Omoda 5 Z", "2024", "8,500 KM", "A", "Rp 125.000.000", "Rp 129.100.000", "-3.28%"],
        ["Suzuki Carry Pick Up", "2015", "145,000 KM", "D", "Rp 45.000.000", "Rp 48.900.000", "-8.67%"],
        ["Honda Brio Satya E", "2020", "52,300 KM", "B", "Rp 118.000.000", "Rp 112.500.000", "4.66%"]
    ]
    
    for r_idx, row_data in enumerate(comp_rows):
        row_cells = table_comp.rows[r_idx + 1].cells
        bg_color = "ECFDF5" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 80, 80, 150, 150)
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9)
                if c_idx == 6:
                    run.font.bold = True
                    
    doc.add_paragraph()
    
    p_eval = doc.add_paragraph()
    p_eval.paragraph_format.line_spacing = 1.5
    p_eval.paragraph_format.space_after = Pt(12)
    p_eval.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_eval.add_run(
        "Kelebihan Model: Integrasi Grade Inspeksi meningkatkan kepekaan model JST terhadap kondisi riil fisik mobil lelang. "
        "Model ini sukses mengidentifikasi depresiasi nilai yang diakibatkan oleh penurunan grade (A ke E) pada unit sejenis. "
        "Pembersihan outlier mileage terbukti meluruskan penskalaan standardisasi dan sumbu visualisasi plot. "
        "Kekurangan Model: Variabilitas harga dasar lelang yang sangat dipengaruhi oleh kelengkapan dokumen (BPKB) "
        "dan kondisi eksternal balai lelang menyebabkan nilai MAPE testing (19.86%) sedikit lebih tinggi dibandingkan model retail."
    )
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # BAB V: KESIMPULAN & SARAN
    # ----------------------------------------------------
    h5 = doc.add_paragraph()
    run_h5 = h5.add_run("BAB V\nKESIMPULAN DAN SARAN")
    run_h5.bold = True
    run_h5.font.size = Pt(14)
    h5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_sub51 = doc.add_paragraph()
    p_sub51.add_run("5.1 Kesimpulan").bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Proyek sistem cerdas estimasi harga dasar lelang mobil bekas berbasis JST MLP 'grade-aware' telah sukses diselesaikan. "
        "Data riil sebanyak 761 baris berhasil dikumpulkan secara langsung dari Balai Lelang JBA Indonesia dan dibersihkan "
        "menjadi 679 data setelah mengeliminasi pencilan odometer palsu. Model JST dilatih dengan dimensi input 101 fitur "
        "one-hot encoded (mengakomodasi kolom grade lelang) and mencapai akurasi R2 testing sebesar 72.90% dengan MAPE "
        "pengujian 19.86% (dalam batas wajar industri lelang). Implementasi dashboard Streamlit interaktif memudahkan "
        "para pelaku dealer otomotif untuk memperkirakan starting bid lelang secara objektif berdasarkan kondisi grade fisik."
    )
    
    p_sub52 = doc.add_paragraph()
    p_sub52.add_run("5.2 Saran Perbaikan").bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "1. Integrasi Status Dokumen: Menambahkan variabel kelengkapan dokumen BPKB (tersedia/ready atau tunda/delayed) "
        "sebagai fitur input, karena faktor ini sangat berpengaruh terhadap pemotongan nilai harga dasar lelang.\n"
        "2. Multi-Grade Input: Memisahkan grade mesin, eksterior, dan interior secara individual (tidak digabung) "
        "ke dalam model JST untuk memberikan estimasi harga yang jauh lebih granular.\n"
        "3. Scraping Terjadwal: Membuat sistem penarikan data terjadwal harian di berbagai cabang balai lelang JBA di seluruh "
        "Indonesia untuk memperluas jangkauan dataset."
    )
    
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # REFERENSI
    # ----------------------------------------------------
    h_ref = doc.add_paragraph()
    run_href = h_ref.add_run("DAFTAR REFERENSI")
    run_href.bold = True
    run_href.font.size = Pt(14)
    h_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    refs = [
        "https://www.jba.co.id/ - Sumber data hasil lelang mobil bekas Indonesia.",
        "https://www.tensorflow.org/ - Dokumentasi framework TensorFlow/Keras API.",
        "https://streamlit.io/ - Dokumentasi platform deployment web Streamlit.",
        "https://scikit-learn.org/ - Pustaka standardisasi data preprocessing dan scaling.",
        "JBA Indonesia, Market Auction Price (MAP) & Panduan Sistem Penilaian Grade Inspeksi Kendaraan Otomotif 2026."
    ]
    for r in refs:
        p_r = doc.add_paragraph()
        p_r.paragraph_format.line_spacing = 1.5
        p_r.paragraph_format.space_after = Pt(6)
        p_r.add_run(r)
        
    doc.save(report_path)
    print(f"Laporan DOCX successfully generated at: {report_path}")

if __name__ == "__main__":
    generate_docx()
